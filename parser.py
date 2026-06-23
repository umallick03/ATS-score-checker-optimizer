import os
import re
from pypdf import PdfReader
from docx import Document

class ResumeParser:
    @staticmethod
    def parse_pdf(file_path: str) -> dict:
        """Parses PDF and extracts text + identifies layout structural issues."""
        text = ""
        structural_issues = []
        try:
            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            
            has_multi_column_spacing = False
            has_non_standard_bullets = False
            
            # Non-standard bullets match common characters like ➢, ❖, ✔, ✓, ▪, ●, etc.
            non_std_bullets_rx = re.compile(
                r'[\u27a2\u2756\u2714\u2713\u25aa\u25cf\u25fe\u25fd\u2b24\u2b1d\u25c6\u25c8\u25cb\u25a0\u25a1\u25aa\u25ab]'
            )
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if not page_text:
                    continue
                text += page_text + "\n"
                
                # Check for large spacing between words on the same line
                # Often indicates multi-column layout or tables in PDFs
                for line in page_text.split('\n'):
                    if re.search(r'\w{2,}\s{4,}\w{2,}', line):
                        has_multi_column_spacing = True
                    if non_std_bullets_rx.search(line):
                        has_non_standard_bullets = True
            
            if has_multi_column_spacing:
                structural_issues.append({
                    "issue": "Multi-column layout or table-like formatting detected",
                    "severity": "Warning",
                    "description": "Multi-column layouts can confuse older ATS scanners, causing them to read text out of order."
                })
            if has_non_standard_bullets:
                structural_issues.append({
                    "issue": "Non-standard bullet points used",
                    "severity": "Warning",
                    "description": "Using decorative shapes (stars, arrows, squares) can disrupt parsing. Stick to standard circular bullets (•) or hyphens (-)."
                })
            
            if len(text.strip()) < 100:
                structural_issues.append({
                    "issue": "Scanned or image-only PDF detected",
                    "severity": "Critical",
                    "description": "The PDF contains very little or no selectable text. Ensure your resume is saved as a text-based PDF rather than a scanned image."
                })
                
            return {
                "text": text,
                "structural_issues": structural_issues,
                "format": "pdf",
                "success": True
            }
        except Exception as e:
            return {
                "error": f"Failed to parse PDF: {str(e)}", 
                "text": "", 
                "structural_issues": [], 
                "format": "pdf",
                "success": False
            }

    @staticmethod
    def parse_docx(file_path: str) -> dict:
        """Parses DOCX and extracts text + identifies layout structural issues."""
        structural_issues = []
        try:
            doc = Document(file_path)
            
            paragraphs_text = []
            for p in doc.paragraphs:
                paragraphs_text.append(p.text)
            
            # Check for tables
            if len(doc.tables) > 0:
                structural_issues.append({
                    "issue": "Tables detected in document",
                    "severity": "Warning",
                    "description": "Tables are frequently unparseable or parsed out of order by ATS. Use simple text layouts instead."
                })
            
            # Check for columns
            has_columns = False
            for section in doc.sections:
                try:
                    cols = section._sectPr.xpath('w:cols')
                    if cols:
                        num_cols = cols[0].get(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'
                        )
                        if num_cols and int(num_cols) > 1:
                            has_columns = True
                except Exception:
                    pass
            
            if has_columns:
                structural_issues.append({
                    "issue": "Multi-column section detected",
                    "severity": "Warning",
                    "description": "Multi-column layouts can confuse older ATS scanners, causing them to read text out of order."
                })
            
            non_std_bullets_rx = re.compile(
                r'[\u27a2\u2756\u2714\u2713\u25aa\u25cf\u25fe\u25fd\u2b24\u2b1d\u25c6\u25c8\u25cb\u25a0\u25a1\u25aa\u25ab]'
            )
            has_non_standard_bullets = False
            
            full_text = "\n".join(paragraphs_text)
            
            # Extract table text to include in full text search
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_texts.append(" | ".join(row_text))
            
            if table_texts:
                full_text += "\n" + "\n".join(table_texts)
                
            for line in full_text.split('\n'):
                if non_std_bullets_rx.search(line):
                    has_non_standard_bullets = True
                    break
                    
            if has_non_standard_bullets:
                structural_issues.append({
                    "issue": "Non-standard bullet points used",
                    "severity": "Warning",
                    "description": "Using decorative shapes (stars, arrows, squares) can disrupt parsing. Stick to standard circular bullets (•) or hyphens (-)."
                })
                
            if len(full_text.strip()) < 100:
                structural_issues.append({
                    "issue": "Empty or extremely short document",
                    "severity": "Critical",
                    "description": "The resume has very little text. Please ensure it is populated with content."
                })
                
            return {
                "text": full_text,
                "structural_issues": structural_issues,
                "format": "docx",
                "success": True
            }
        except Exception as e:
            return {
                "error": f"Failed to parse DOCX: {str(e)}", 
                "text": "", 
                "structural_issues": [], 
                "format": "docx",
                "success": False
            }

    @classmethod
    def parse_file(cls, file_path: str) -> dict:
        _, ext = os.path.splitext(file_path.lower())
        if ext == '.pdf':
            return cls.parse_pdf(file_path)
        elif ext == '.docx':
            return cls.parse_docx(file_path)
        else:
            return {
                "error": f"Unsupported file format: {ext}. Only PDF and DOCX are supported.",
                "text": "",
                "structural_issues": [],
                "format": ext.replace('.', ''),
                "success": False
            }
